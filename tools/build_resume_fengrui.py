# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parents[1] / "outputs"
OUT_FILE = OUT_DIR / "冯国睿_AI应用开发工程师_新版简历.docx"

BLUE = RGBColor(31, 77, 120)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(86, 96, 112)
INK = RGBColor(20, 28, 38)
BORDER = "D9E2EC"
LIGHT = "F4F7FA"


def set_run_font(run, size=10.0, bold=False, color=INK, italic=False, ascii_font="Calibri", east_font="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_para(paragraph, before=0, after=3, line=1.08, left=None, first=None, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if left is not None:
        fmt.left_indent = left
    if first is not None:
        fmt.first_line_indent = first
    if align is not None:
        paragraph.alignment = align


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_para(p, before=8, after=4, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=BLUE)
    paragraph_bottom_border(p, color=RGBColor(217, 226, 236), size="6")
    return p


def add_plain(doc, text, size=9.6, after=2, bold=False, color=INK):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.08)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, size=9.4, after=1.5):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, after=after, line=1.08, left=Inches(0.23), first=Inches(-0.14))
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK)
    return p


def add_project(doc, title, period, stack, intro, bullets):
    p = doc.add_paragraph()
    set_para(p, before=5, after=1, line=1.0)
    r = p.add_run(title)
    set_run_font(r, size=10.4, bold=True, color=INK)
    r = p.add_run(f"  |  {period}")
    set_run_font(r, size=9.2, color=MUTED)

    p = doc.add_paragraph()
    set_para(p, before=0, after=1.5, line=1.0)
    r = p.add_run("技术栈：")
    set_run_font(r, size=9.0, bold=True, color=BLUE)
    r = p.add_run(stack)
    set_run_font(r, size=9.0, color=MUTED)

    add_plain(doc, intro, size=9.3, after=1.5, color=INK)
    for item in bullets:
        add_bullet(doc, item)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.6)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(9.3)
    bullet.paragraph_format.space_after = Pt(1.5)
    bullet.paragraph_format.line_spacing = 1.08

    # Header
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [5000, 4480])
    for cell in table.row_cells(0):
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell, color="FFFFFF", size="0")
    left, right = table.row_cells(0)

    p = left.paragraphs[0]
    set_para(p, after=0, line=1.0)
    r = p.add_run("冯国睿")
    set_run_font(r, size=21, bold=True, color=INK)
    p = left.add_paragraph()
    set_para(p, after=0, line=1.0)
    r = p.add_run("AI 应用开发工程师  |  RAG / Agent / AI 产品工程")
    set_run_font(r, size=10.4, bold=True, color=BLUE)

    contacts = [
        "手机：18746115091",
        "邮箱：2731025858@qq.com",
        "GitHub：https://github.com/hhh123-ffff?tab=repositories",
        "22岁 | 男 | 一周内到岗",
    ]
    for idx, item in enumerate(contacts):
        p = right.paragraphs[0] if idx == 0 else right.add_paragraph()
        set_para(p, after=0.5, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run(item)
        set_run_font(r, size=8.8, color=MUTED)

    add_plain(
        doc,
        "数据科学与大数据技术本科背景，聚焦大模型应用落地。具备 RAG 知识库、AI Agent 工作流、AI 译制工作台等项目实践，能够完成从需求拆解、接口与数据模型设计、模型/第三方服务接入到调试优化的完整开发闭环。",
        size=9.7,
        after=3,
    )

    add_heading(doc, "教育背景")
    edu = doc.add_table(rows=1, cols=3)
    set_table_width(edu, [2250, 4500, 2730])
    for cell in edu.row_cells(0):
        set_cell_border(cell, color=BORDER, size="4")
    vals = ["2022.09 - 2026.06", "北京交通大学海滨学院 | 数据科学与大数据技术 | 本科", "团支书"]
    for cell, val in zip(edu.row_cells(0), vals):
        p = cell.paragraphs[0]
        set_para(p, after=0, line=1.05)
        r = p.add_run(val)
        set_run_font(r, size=9.3, bold=(val == vals[1]), color=INK if val != vals[0] else MUTED)
    add_plain(doc, "相关课程：Python 程序设计、数据库原理、数据可视化、概率论与数理统计、线性代数、NLP 基础。", size=9.0, after=1.5, color=MUTED)

    add_heading(doc, "核心技能")
    skill_rows = [
        ("AI 大模型 / RAG", "Prompt Engineering、Function Calling、Embedding、向量检索、Top-K 上下文构建、FAISS / Chroma、LangChain / LangGraph、Ollama、本地模型与多模型切换。"),
        ("后端与数据", "Python、Java、TypeScript / JavaScript、FastAPI、Express、RESTful API、Zod、PostgreSQL / MySQL、Pandas、数据清洗与结构化处理。"),
        ("Agent 与工作流", "AI Agent 工具调用编排、StateGraph、DAG 调度、节点抽象、变量映射、SSE / 任务轮询、Coze / Dify 工作流设计理解。"),
        ("工程化能力", "Git、Docker、Linux 基础部署、S3 / Local 文件存储、Cookie Session、Argon2、RBAC / capability 权限、限流、审计日志、Vitest。"),
    ]
    skills = doc.add_table(rows=len(skill_rows), cols=2)
    set_table_width(skills, [1800, 7680])
    for row_idx, (label, detail) in enumerate(skill_rows):
        cells = skills.row_cells(row_idx)
        for cell in cells:
            set_cell_border(cell, color=BORDER, size="4")
        set_cell_shading(cells[0], LIGHT)
        p = cells[0].paragraphs[0]
        set_para(p, after=0, line=1.05)
        r = p.add_run(label)
        set_run_font(r, size=9.1, bold=True, color=BLUE)
        p = cells[1].paragraphs[0]
        set_para(p, after=0, line=1.05)
        r = p.add_run(detail)
        set_run_font(r, size=9.0, color=INK)

    add_heading(doc, "项目经历")
    add_project(
        doc,
        "聆溪智译：影视短剧 AI 译制工作台",
        "2026.04 - 2026.05",
        "Next.js / React / TypeScript / Express / PostgreSQL / Zod / Volcengine VOD / ElevenLabs / S3",
        "面向短剧译制场景构建 AI 工作台，覆盖登录、剧集与分集、视频上传、自动译制、字幕编辑、角色音色、配音生成、音频时间线、导出、团队协作与会员积分等核心链路。",
        [
            "梳理并实现前后端分离架构：Next.js 前端通过 /api 代理访问独立 Express API，后端按 auth、workspace、series、episode、translation、dubbing、voice 等领域拆分路由、服务与仓储。",
            "设计 PostgreSQL 数据模型，覆盖 27 张业务表与 13 个迁移文件，支持多租户、素材/翻译/配音版本化、统一任务状态、审计日志与账务积分等能力。",
            "建设页面聚合读模型与工作台核心链路，支撑字幕编辑、SRT 导入导出、单条/批量配音、音色选择、情绪复核、音频片段时间线保存等高频操作。",
            "抽象第三方 Provider Ports，接入 Volcengine 视频上传/自动译制与 ElevenLabs 音色目录、音色克隆、文本转语音能力，并支持 mock / 真实供应商按环境切换。",
            "完善认证与权限基础设施：Cookie Session、Argon2id 密码、workspace role/capability matrix、CORS allowlist、接口限流、关键写操作审计。",
            "沉淀 22 个后端测试文件，覆盖认证、账务、配音、Provider、读模型、SRT 导入导出、权限矩阵等核心模块，降低复杂工作台迭代风险。",
        ],
    )

    add_project(
        doc,
        "基于 RAG 的企业级智能知识库问答系统",
        "2026.01 - 2026.03",
        "Python / FastAPI / Chroma / Embedding / Ollama / Prompt Engineering / Streaming Response",
        "面向企业制度、产品资料与 FAQ 的智能问答场景，实现文档解析、向量化索引、语义检索与生成式回答闭环，提升内部知识获取效率。",
        [
            "基于 FastAPI 设计文档上传、文本切分、向量化、检索问答等接口，完成从数据入库到问答生成的 RAG Pipeline。",
            "使用 Chroma 构建向量知识库，结合 Top-K 检索、上下文拼接与 Prompt 模板，将召回内容结构化注入模型输入。",
            "优化 chunk size / overlap、检索条数与上下文构建策略，改善召回不准和上下文不足问题，典型问答相关性较初版提升约 30%+。",
            "接入 Ollama 本地大模型并封装模型调用/切换能力，支持私有化部署；通过来源约束、结构化输出和去重逻辑降低幻觉与冗余回答。",
            "支持基础流式输出，文档检索响应时间控制在 1-2 秒内，满足常见内部知识库交互需求。",
        ],
    )

    add_project(
        doc,
        "基于 LangGraph 的企业级 AI 工作流编排平台",
        "2026.03 - 2026.04",
        "LangGraph / StateGraph / FastAPI / SSE / OpenAI / DeepSeek / 通义千问 / Tool Calling",
        "面向企业 AI Agent 构建与任务自动化场景，设计可编排的工作流执行引擎，支持多模型接入、工具节点组合、条件分支与运行状态回传。",
        [
            "基于 LangGraph StateGraph 构建工作流执行引擎，实现节点注册、边连接、状态流转、条件分支和图结构任务编排。",
            "封装模型工厂与配置管理模块，支持 OpenAI、DeepSeek、通义千问等模型动态切换与参数配置。",
            "抽象通用节点执行框架，统一处理 Prompt 渲染、参数解析、模型调用、工具执行和结果输出，提升节点扩展能力。",
            "实现变量解析与上下文映射机制，支持节点间数据传递、运行时动态参数注入和工具调用结果复用。",
            "基于 DAG 拓扑排序控制执行顺序，加入依赖解析与循环依赖检测；结合 FastAPI + SSE 回传执行状态与流式输出。",
        ],
    )

    add_heading(doc, "个人优势")
    for item in [
        "学习能力强，能快速进入新技术领域，并通过项目实践将 RAG、Agent、AI 工作流和第三方 AI 服务落到可运行系统中。",
        "做事认真主动，遇到召回不准、上下文不足、模型幻觉、复杂页面状态和第三方接入等问题，能够持续定位、拆解并优化。",
        "具备从需求分析、架构设计、接口实现、数据建模、测试验证到部署联调的完整闭环意识，适合 AI 应用开发与大模型产品工程岗位。",
    ]:
        add_bullet(doc, item, size=9.3, after=1.5)

    footer = section.footer.paragraphs[0]
    set_para(footer, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = footer.add_run("冯国睿 | AI 应用开发工程师")
    set_run_font(r, size=8.0, color=MUTED)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_doc())
